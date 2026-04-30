from typing import Optional
import os
import re
import sys
import copy
import subprocess
import pandas as pd
from tkinter import filedialog
from docx import Document
import win32com.client
import pythoncom


def kill_winword_processes(reason: str = ""):
    """强制结束所有 WINWORD.EXE 进程，避免 COM 附着到挂着隐藏弹窗的僵尸 Word。"""
    tag = f"（{reason}）" if reason else ""
    print(f"🧹 正在清理残留 WINWORD.EXE 进程{tag}...")
    try:
        result = subprocess.run(
            ["taskkill", "/F", "/IM", "WINWORD.EXE", "/T"],
            capture_output=True, text=True, timeout=10,
        )
        if result.returncode == 0:
            print("   ↳ 已结束残留 Word 进程")
        else:
            # returncode 128 / 1 通常代表"没有匹配进程"，正常
            print("   ↳ 没有发现需要清理的 Word 进程")
    except Exception as e:
        print(f"   ⚠️ taskkill 调用失败（忽略，继续）: {e}")


def unblock_file(file_path: str):
    """移除 Windows 的"来自互联网"标记 (Zone.Identifier ADS)，避免 Word 触发受保护视图。"""
    abs_path = os.path.abspath(file_path)
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", f"Unblock-File -LiteralPath \"{abs_path}\""],
            capture_output=True, text=True, timeout=10,
        )
        if result.returncode == 0:
            print(f"🔓 已解除文件网络标记: {os.path.basename(abs_path)}")
        else:
            print(f"   ⚠️ Unblock-File 返回非 0（通常无标记，可忽略）: {result.stderr.strip()}")
    except Exception as e:
        print(f"   ⚠️ Unblock-File 调用失败（忽略，继续）: {e}")

# 让 print 实时刷新，避免在 IDE / 重定向场景下日志被块缓冲攒到最后才一次性输出
_reconfigure = getattr(sys.stdout, "reconfigure", None)
if callable(_reconfigure):
    try:
        _reconfigure(line_buffering=True)
    except Exception:
        pass

# 【架构解耦】：直接从你的通用 UI 库中引入动态表单组件
from ui_components import ModernDynamicFormDialog

# ==========================================
# 模块 1：配置中心
# ==========================================
class Config:
    EXCEL_PATH: str = '缺陷清单.xlsx'
    EXCEL_SHEET: Optional[str] = None  # None = 第一个 sheet
    EXCEL_COL_NAME: str = '照片'
    WORD_PATH: str = '待排序_附录1.docx'
    OUTPUT_DIR: str = ''  # 空 = 与 Word 同目录
    OUTPUT_NAME: str = '已排序_附录1.docx'
    OUTPUT_PATH: str = '已排序_附录1.docx'  # 由 OUTPUT_DIR + OUTPUT_NAME 拼出
    MATCH_PATTERN: str = r'图\s*(\d+)'

    @classmethod
    def update(cls, **kwargs):
        for k, v in kwargs.items():
            setattr(cls, k, v)

# ==========================================
# 模块 2：解析器（提取排序规则）
# ==========================================
def get_excel_sort_order(path, col_name, sheet_name=None):
    try:
        df = pd.read_excel(path, sheet_name=sheet_name) if sheet_name else pd.read_excel(path)
        if col_name not in df.columns:
            print(f"❌ Sheet [{sheet_name or '默认'}] 中找不到列：{col_name}")
            return []
        raw_list = df[col_name].dropna().astype(str).tolist()
        order = []
        for item in raw_list:
            match = re.search(Config.MATCH_PATTERN, item)
            if match:
                order.append(int(match.group(1)))
        return order
    except Exception as e:
        print(f"❌ 读取 Excel 失败: {e}")
        return []

# ==========================================
# 模块 3：文档重构器（核心排序逻辑）
# ==========================================
class WordTableProcessor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.pairs = {}
        self.unmatched = []

    def extract_pairs(self, excel_order):
        """使用 python-docx 提取配对信息（仅文本，用于匹配）"""
        print("🔍 正在用 python-docx 解析 Word 文档...")
        doc = Document(self.doc_path)

        if not doc.tables:
            raise ValueError("❌ Word 文档中没有找到任何表格！")

        table = doc.tables[0]
        total_rows = len(table.rows)
        print(f"📑 检测到表格共 {total_rows} 行，开始扫描配对（含合并单元格时此步可能较慢）...")

        excel_order_set = set(excel_order)  # 加速 in 判断

        for i in range(0, total_rows, 2):
            if i + 1 >= total_rows:
                break

            if i % 20 == 0:
                print(f"   ↳ 解析进度: {i}/{total_rows}")

            img_row = table.rows[i]
            txt_row = table.rows[i+1]

            for j in range(len(img_row.cells)):
                txt_cell = txt_row.cells[j]
                text = txt_cell.text.strip()
                if not text:
                    continue

                match = re.search(Config.MATCH_PATTERN, text)
                if match:
                    num = int(match.group(1))
                    # 记录行索引，用于后续 COM 处理
                    pair_data = {"num": num, "img_row_idx": i, "txt_row_idx": i+1, "img_col_idx": j, "txt_col_idx": j}

                    if num in excel_order_set:
                        self.pairs[num] = pair_data
                    else:
                        self.unmatched.append(pair_data)

        print(f"✅ 解析完成：匹配 {len(self.pairs)} 个，未匹配 {len(self.unmatched)} 个")

    def rebuild(self, excel_order):
        """使用 COM 接口正确复制包含图片的单元格"""
        # 启动前先把所有 WINWORD.EXE 干掉，确保我们启动的是一个干净进程
        kill_winword_processes(reason="启动 Word COM 前预清理")

        # 初始化 COM
        pythoncom.CoInitialize()
        word_app = None

        try:
            # 先解除"来自互联网"标记，避免触发受保护视图导致 Open 卡住
            unblock_file(self.doc_path)

            print("📋 正在启动 Word 应用（DispatchEx 强制新进程；窗口设为可见，方便看到任何弹框）...")
            word_app = win32com.client.DispatchEx("Word.Application")
            word_app.Visible = True  # ← 暂时设为 True 用于调试；问题定位后可改回 False
            word_app.DisplayAlerts = 0
            # 关闭可能导致隐藏弹窗的功能
            try:
                word_app.AutomationSecurity = 3  # msoAutomationSecurityForceDisable
            except Exception:
                pass

            # 打开源文档（显式传所有可能弹框的开关，强制静默打开）
            print("📂 正在打开源文档...")
            src_path_abs = os.path.abspath(self.doc_path)
            src_doc = word_app.Documents.Open(
                FileName=src_path_abs,
                ConfirmConversions=False,   # 不弹"格式转换"对话框
                ReadOnly=False,
                AddToRecentFiles=False,
                Revert=False,               # 已打开则不重新加载
                Format=0,                   # wdOpenFormatAuto
                Visible=True,
                OpenAndRepair=False,        # 不触发"打开并修复"流程，那个会卡很久
                NoEncodingDialog=True,
            )
            print("✅ 源文档已打开")
            
            # 创建新文档
            print("📄 正在创建新文档...")
            new_doc = word_app.Documents.Add()
            
            if new_doc.Tables.Count > 0:
                new_doc.Tables(1).Delete()
            
            # 获取源表格
            src_table = src_doc.Tables(1)
            src_rows = src_table.Rows.Count
            src_cols = src_table.Columns.Count
            print(f"📊 源表格共有 {src_rows} 行 × {src_cols} 列")

            # 构建最终顺序列表
            final_list = []
            for num in excel_order:
                if num in self.pairs:
                    final_list.append(self.pairs[num])
            final_list.extend(self.unmatched)
            print(f"📋 共有 {len(final_list)} 个配对待处理")

            # 还原源表的版式：每"组"占 2 行（图片行 + 说明行）× src_cols 列
            # 即每行能容纳 src_cols 张图，前后两行成对（上图下注）
            pairs_per_group = src_cols
            num_groups = (len(final_list) + pairs_per_group - 1) // pairs_per_group
            total_new_rows = num_groups * 2

            print(f"🔨 正在创建新表格 ({total_new_rows} 行 × {src_cols} 列，{pairs_per_group} 张图/组)...")
            new_table = new_doc.Tables.Add(
                Range=new_doc.Range(0, 0),
                NumRows=total_new_rows,
                NumColumns=src_cols,
                DefaultTableBehavior=1  # wdWord9TableBehavior
            )
            new_table.Borders.Enable = True
            print("✅ 新表格创建完成")

            # 逐个复制单元格：第 idx 个配对 → group=idx//cols，col=idx%cols+1，图行=group*2+1，注行=group*2+2
            for idx, item in enumerate(final_list):
                if idx % 10 == 0:
                    print(f"🔄 进度: {idx+1}/{len(final_list)}")

                group_idx = idx // pairs_per_group
                col_in_new = (idx % pairs_per_group) + 1
                img_row_in_new = group_idx * 2 + 1
                txt_row_in_new = group_idx * 2 + 2

                # 复制图片单元格 → 新表的"图行"
                src_img_cell = src_table.Cell(Row=item["img_row_idx"] + 1, Column=item["img_col_idx"] + 1)
                target_img_cell = new_table.Cell(Row=img_row_in_new, Column=col_in_new)
                try:
                    src_img_cell.Range.Copy()
                    target_img_cell.Range.Paste()
                except Exception as e:
                    print(f"⚠️ 复制图片单元格 ({idx}) 失败: {e}")
                    target_img_cell.Range.Text = "[图片]"

                # 复制文字单元格 → 新表的"注行"
                src_txt_cell = src_table.Cell(Row=item["txt_row_idx"] + 1, Column=item["txt_col_idx"] + 1)
                target_txt_cell = new_table.Cell(Row=txt_row_in_new, Column=col_in_new)
                try:
                    src_txt_cell.Range.Copy()
                    target_txt_cell.Range.Paste()
                except Exception as e:
                    print(f"⚠️ 复制文字单元格 ({idx}) 失败: {e}")
                    target_txt_cell.Range.Text = src_txt_cell.Range.Text

            # 末尾若不满一组（例如 src_cols=2 但配对数为奇数），剩余单元格留空即可，无需特殊处理
            
            # 保存新文档
            print("💾 正在保存新文档...")
            new_doc.SaveAs(os.path.abspath(Config.OUTPUT_PATH))
            print("✅ 文档已保存")
            
            new_doc.Close()
            src_doc.Close()
            
            print(f"✅ 处理完成！已生成新文件: {Config.OUTPUT_PATH}")
            
        except Exception as e:
            print(f"❌ COM 重构失败: {e}")
            raise
        finally:
            if word_app:
                try:
                    word_app.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()
            # 兜底清理：即使 Quit 没释放干净，也确保不会有僵尸进程留下来卡住下次运行
            kill_winword_processes(reason="退出兜底")

# ==========================================
# 模块 4：执行入口（纯粹的业务调用）
# ==========================================
def pick_excel_first():
    """先单独弹一个原生文件框选 Excel，目的是为了能预读 Sheet 列表"""
    path = filedialog.askopenfilename(
        title="第一步：选择缺陷清单 Excel",
        filetypes=[("Excel 文件", "*.xlsx *.xlsm *.xls"), ("所有文件", "*.*")]
    )
    return path

def read_sheet_names(excel_path):
    try:
        return pd.ExcelFile(excel_path).sheet_names
    except Exception as e:
        print(f"❌ 读取 Sheet 列表失败: {e}")
        return []

def request_sort_params(excel_path, sheet_names):
    """弹出主参数窗口，让用户挑选 Sheet、列名、Word、输出位置"""
    default_dir = os.path.dirname(excel_path) or os.getcwd()
    schema = [
        {"key": "sheet_name", "label": "Excel 工作表:", "type": "select",
         "options": sheet_names, "default": sheet_names[0] if sheet_names else ""},
        {"key": "excel_col", "label": "照片列表头:", "type": "text",
         "default": Config.EXCEL_COL_NAME},
        {"key": "word_path", "label": "待排序 Word:", "type": "file",
         "file_types": [("Word 文件", "*.docx"), ("所有文件", "*.*")],
         "default": Config.WORD_PATH},
        {"key": "output_dir", "label": "输出目录:", "type": "dir",
         "default": default_dir},
        {"key": "output_name", "label": "输出文件名:", "type": "text",
         "default": Config.OUTPUT_NAME},
    ]
    dialog = ModernDynamicFormDialog(title="照片排序 - 参数配置", form_schema=schema, width=620)
    return dialog.show()


if __name__ == "__main__":
    # 阶段一：先选 Excel 文件，才能读出可选的 Sheet 列表
    excel_path = pick_excel_first()
    if not excel_path:
        print("⚠️ 已取消：未选择 Excel 文件。")
        raise SystemExit

    sheets = read_sheet_names(excel_path)
    if not sheets:
        print("⚠️ 终止：该 Excel 没有可读取的工作表。")
        raise SystemExit

    # 阶段二：动态表单收集其余参数（Sheet 下拉框 + 输出位置）
    params = request_sort_params(excel_path, sheets)

    if not params or not params.get("word_path"):
        print("⚠️ 已取消：未选择 Word 文件或直接关闭了窗口。")
    else:
        Config.update(
            EXCEL_PATH=excel_path,
            EXCEL_SHEET=params.get("sheet_name") or sheets[0],
            EXCEL_COL_NAME=params.get("excel_col") or Config.EXCEL_COL_NAME,
            WORD_PATH=params["word_path"],
            OUTPUT_DIR=params.get("output_dir") or os.path.dirname(Config.WORD_PATH),
            OUTPUT_NAME=params.get("output_name") or Config.OUTPUT_NAME
        )

        # 文件名安全兜底：如果用户没写后缀，自动补 .docx
        if not Config.OUTPUT_NAME.lower().endswith(".docx"):
            Config.OUTPUT_NAME += ".docx"

        Config.OUTPUT_PATH = os.path.join(Config.OUTPUT_DIR, Config.OUTPUT_NAME)

        print(f"🚀 正在读取数据... [Sheet: {Config.EXCEL_SHEET}]")
        order = get_excel_sort_order(Config.EXCEL_PATH, Config.EXCEL_COL_NAME, Config.EXCEL_SHEET)

        if not order:
            print("⚠️ 终止：Excel 排序数据为空。")
        else:
            print(f"📊 从 Excel 获取到 {len(order)} 个排序指令，开始重构 Word...")
            processor = WordTableProcessor(Config.WORD_PATH)
            print("—— 阶段 A：解析 Word 表格 ——")
            processor.extract_pairs(order)
            print("—— 阶段 B：通过 Word COM 重构文档 ——")
            processor.rebuild(order)
