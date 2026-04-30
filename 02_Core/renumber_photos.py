"""
照片重编号工具 —— sort_photos.py 的下游配套工具

工作流：sort_photos.py 把照片按 Excel 缺陷清单顺序排好后，每张照片的题注还是原编号（图17、图3、图25…），
本程序按"行优先"（左→右、上→下）扫描已排序 Word 表格，把题注重命名为 图1、图2、图3…，
并同步更新 Excel 缺陷清单的"照片"列里的所有"图 N"引用。
"""
from typing import Optional, Dict
import os
import re
import sys
from tkinter import filedialog
from docx import Document
from openpyxl import load_workbook
import pandas as pd

# 让 print 实时刷新（与 sort_photos.py 一致）
_reconfigure = getattr(sys.stdout, "reconfigure", None)
if callable(_reconfigure):
    try:
        _reconfigure(line_buffering=True)
    except Exception:
        pass

from ui_components import ModernDynamicFormDialog


# ==========================================
# 模块 1：配置中心
# ==========================================
class Config:
    EXCEL_PATH: str = '缺陷清单.xlsx'
    EXCEL_SHEET: Optional[str] = None
    EXCEL_COL_NAME: str = '照片'
    WORD_PATH: str = '已排序_附录1.docx'
    OUTPUT_DIR: str = ''
    OUTPUT_WORD_NAME: str = '已重编号_附录1.docx'
    OUTPUT_EXCEL_NAME: str = '已重编号_缺陷清单.xlsx'
    OUTPUT_WORD_PATH: str = ''
    OUTPUT_EXCEL_PATH: str = ''
    MATCH_PATTERN: str = r'图\s*(\d+)'

    @classmethod
    def update(cls, **kwargs):
        for k, v in kwargs.items():
            setattr(cls, k, v)


# ==========================================
# 模块 2：映射构建器（扫描 Word 表格生成 旧→新 编号映射）
# ==========================================
def build_renumber_mapping(doc_path: str) -> Dict[int, int]:
    """按"行优先"顺序遍历已排序 Word 文档第一个表格的题注行，
    返回 {旧编号: 新编号}，新编号从 1 开始递增。

    源表结构（由 sort_photos.py 生成）：
      第 0 行：图 | 图 | …
      第 1 行：注 | 注 | …   ← 题注在这里
      第 2 行：图 | 图 | …
      第 3 行：注 | 注 | …
      …
    """
    print(f"🔍 扫描 Word 文档构建编号映射: {os.path.basename(doc_path)}")
    doc = Document(doc_path)
    if not doc.tables:
        raise ValueError("❌ Word 文档中没有找到任何表格！")

    table = doc.tables[0]
    pattern = re.compile(Config.MATCH_PATTERN)
    mapping: Dict[int, int] = {}
    new_num = 1

    # 题注行 = 0-indexed 的奇数行 (1, 3, 5, …)
    for row_idx in range(1, len(table.rows), 2):
        row = table.rows[row_idx]
        seen_tc = set()  # 去重合并单元格
        # 行内左→右：python-docx 的 row.cells 已经按列序返回
        for cell in row.cells:
            if cell._tc in seen_tc:
                continue
            seen_tc.add(cell._tc)

            text = cell.text.strip()
            if not text:
                continue

            match = pattern.search(text)
            if not match:
                continue

            old_num = int(match.group(1))
            if old_num in mapping:
                # 同一编号在多个题注里出现，理论上不该发生 —— 给个提示但不中断
                print(f"   ⚠️ 编号 {old_num} 重复出现，已忽略后续映射")
                continue
            mapping[old_num] = new_num
            new_num += 1

    print(f"✅ 已构建 {len(mapping)} 条编号映射 (1 → {new_num - 1})")
    # 打印前几条用于人工核对
    preview = list(mapping.items())[:8]
    print("   预览（前 8 条 旧→新）:", "  ".join(f"{o}→{n}" for o, n in preview))
    return mapping


# ==========================================
# 模块 3：替换器（按映射改 Word 题注 + Excel 引用）
# ==========================================
def _make_substitutor(mapping: Dict[int, int]):
    """生成一个 re.sub 用的回调，保留'图'和空白前缀，只换数字。"""
    pattern = re.compile(Config.MATCH_PATTERN)
    unmatched = []  # 收集找不到映射的旧编号

    def sub(m: re.Match) -> str:
        old = int(m.group(1))
        new = mapping.get(old)
        if new is None:
            unmatched.append(old)
            return m.group(0)
        # 保留 "图" + 中间空白，只替换数字部分
        prefix = m.group(0)[:m.start(1) - m.start(0)]
        return f"{prefix}{new}"

    def apply(text: str) -> str:
        return pattern.sub(sub, text)

    return apply, unmatched


def renumber_word(doc_path: str, mapping: Dict[int, int], output_path: str):
    """改写 Word 第一个表格的题注，保存为新文档（不覆盖原文件）。"""
    print(f"📝 改写 Word 题注 → {os.path.basename(output_path)}")
    doc = Document(doc_path)
    if not doc.tables:
        raise ValueError("❌ Word 文档中没有找到任何表格！")

    table = doc.tables[0]
    apply, unmatched = _make_substitutor(mapping)
    pattern = re.compile(Config.MATCH_PATTERN)

    replaced_runs = 0
    fallback_paragraphs = 0

    for row_idx in range(1, len(table.rows), 2):
        row = table.rows[row_idx]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # 优先逐 run 替换（保留每段的格式）
                touched = False
                for run in paragraph.runs:
                    if pattern.search(run.text):
                        new_text = apply(run.text)
                        if new_text != run.text:
                            run.text = new_text
                            replaced_runs += 1
                            touched = True

                # 若 run 内没匹配但段落整体有（说明"图 N"被拆到多个 run 里），
                # 退回段落级替换：把替换后的全文塞进第一个 run，其余清空
                if not touched and pattern.search(paragraph.text):
                    full_new = apply(paragraph.text)
                    if full_new != paragraph.text and paragraph.runs:
                        paragraph.runs[0].text = full_new
                        for r in paragraph.runs[1:]:
                            r.text = ""
                        fallback_paragraphs += 1

    doc.save(output_path)
    print(f"   ↳ run 级替换 {replaced_runs} 处，段落级回退 {fallback_paragraphs} 处")
    if unmatched:
        uniq = sorted(set(unmatched))
        print(f"   ⚠️ Word 中有 {len(uniq)} 个编号没有映射（已保留原值）: {uniq[:10]}{' …' if len(uniq) > 10 else ''}")


def renumber_excel(excel_path: str, sheet_name: Optional[str], col_name: str,
                   mapping: Dict[int, int], output_path: str):
    """用 openpyxl 直改 Excel 指定 sheet 的指定列，保留原工作簿其余格式。"""
    print(f"📊 改写 Excel 引用 → {os.path.basename(output_path)}")
    wb = load_workbook(excel_path)

    target_sheet = sheet_name if sheet_name and sheet_name in wb.sheetnames else wb.sheetnames[0]
    ws = wb[target_sheet]

    # 定位"照片"列（默认认为第 1 行是表头）
    header_row = 1
    col_idx = None
    for cell in ws[header_row]:
        if cell.value is not None and str(cell.value).strip() == col_name:
            col_idx = cell.column
            break
    if col_idx is None:
        raise ValueError(f"❌ Sheet [{target_sheet}] 中找不到列：{col_name}")

    apply, unmatched = _make_substitutor(mapping)
    pattern = re.compile(Config.MATCH_PATTERN)

    replaced = 0
    for row in ws.iter_rows(min_row=header_row + 1, min_col=col_idx, max_col=col_idx):
        cell = row[0]
        if cell.value is None:
            continue
        s = str(cell.value)
        if not pattern.search(s):
            continue
        new_s = apply(s)
        if new_s != s:
            cell.value = new_s
            replaced += 1

    wb.save(output_path)
    print(f"   ↳ 已更新 {replaced} 个单元格")
    if unmatched:
        uniq = sorted(set(unmatched))
        print(f"   ⚠️ Excel 中有 {len(uniq)} 个编号没有映射（已保留原值）: {uniq[:10]}{' …' if len(uniq) > 10 else ''}")


# ==========================================
# 模块 4：执行入口（与 sort_photos.py 同样的 UI 流程）
# ==========================================
def pick_excel_first():
    return filedialog.askopenfilename(
        title="第一步：选择已排序的缺陷清单 Excel",
        filetypes=[("Excel 文件", "*.xlsx *.xlsm *.xls"), ("所有文件", "*.*")]
    )


def read_sheet_names(excel_path):
    try:
        return pd.ExcelFile(excel_path).sheet_names
    except Exception as e:
        print(f"❌ 读取 Sheet 列表失败: {e}")
        return []


def request_renumber_params(excel_path, sheet_names):
    default_dir = os.path.dirname(excel_path) or os.getcwd()
    schema = [
        {"key": "sheet_name", "label": "Excel 工作表:", "type": "select",
         "options": sheet_names, "default": sheet_names[0] if sheet_names else ""},
        {"key": "excel_col", "label": "照片列表头:", "type": "text",
         "default": Config.EXCEL_COL_NAME},
        {"key": "word_path", "label": "已排序的 Word:", "type": "file",
         "file_types": [("Word 文件", "*.docx"), ("所有文件", "*.*")],
         "default": Config.WORD_PATH},
        {"key": "output_dir", "label": "输出目录:", "type": "dir",
         "default": default_dir},
        {"key": "output_word_name", "label": "输出 Word 文件名:", "type": "text",
         "default": Config.OUTPUT_WORD_NAME},
        {"key": "output_excel_name", "label": "输出 Excel 文件名:", "type": "text",
         "default": Config.OUTPUT_EXCEL_NAME},
    ]
    dialog = ModernDynamicFormDialog(title="照片重编号 - 参数配置", form_schema=schema, width=620)
    return dialog.show()


if __name__ == "__main__":
    # 阶段一：选 Excel，读 sheet 列表
    excel_path = pick_excel_first()
    if not excel_path:
        print("⚠️ 已取消：未选择 Excel 文件。")
        raise SystemExit

    sheets = read_sheet_names(excel_path)
    if not sheets:
        print("⚠️ 终止：该 Excel 没有可读取的工作表。")
        raise SystemExit

    # 阶段二：动态表单收集其余参数
    params = request_renumber_params(excel_path, sheets)
    if not params or not params.get("word_path"):
        print("⚠️ 已取消：未选择 Word 文件或直接关闭了窗口。")
        raise SystemExit

    Config.update(
        EXCEL_PATH=excel_path,
        EXCEL_SHEET=params.get("sheet_name") or sheets[0],
        EXCEL_COL_NAME=params.get("excel_col") or Config.EXCEL_COL_NAME,
        WORD_PATH=params["word_path"],
        OUTPUT_DIR=params.get("output_dir") or os.path.dirname(params["word_path"]),
        OUTPUT_WORD_NAME=params.get("output_word_name") or Config.OUTPUT_WORD_NAME,
        OUTPUT_EXCEL_NAME=params.get("output_excel_name") or Config.OUTPUT_EXCEL_NAME,
    )

    # 文件名安全兜底
    if not Config.OUTPUT_WORD_NAME.lower().endswith(".docx"):
        Config.OUTPUT_WORD_NAME += ".docx"
    if not Config.OUTPUT_EXCEL_NAME.lower().endswith((".xlsx", ".xlsm")):
        Config.OUTPUT_EXCEL_NAME += ".xlsx"

    Config.OUTPUT_WORD_PATH = os.path.join(Config.OUTPUT_DIR, Config.OUTPUT_WORD_NAME)
    Config.OUTPUT_EXCEL_PATH = os.path.join(Config.OUTPUT_DIR, Config.OUTPUT_EXCEL_NAME)

    print("—— 阶段 A：扫描 Word 构建编号映射 ——")
    mapping = build_renumber_mapping(Config.WORD_PATH)
    if not mapping:
        print("⚠️ 终止：未在 Word 题注中检测到任何'图 N'。")
        raise SystemExit

    print("—— 阶段 B：改写 Word 题注 ——")
    renumber_word(Config.WORD_PATH, mapping, Config.OUTPUT_WORD_PATH)

    print("—— 阶段 C：同步改写 Excel 引用 ——")
    renumber_excel(Config.EXCEL_PATH, Config.EXCEL_SHEET, Config.EXCEL_COL_NAME,
                   mapping, Config.OUTPUT_EXCEL_PATH)

    print(f"\n🎉 全部完成！")
    print(f"   📄 Word: {Config.OUTPUT_WORD_PATH}")
    print(f"   📊 Excel: {Config.OUTPUT_EXCEL_PATH}")
